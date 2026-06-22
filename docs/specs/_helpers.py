"""Shared helpers for Umang HIMS spec document generation.

Generates clean, professional .docx files with consistent styling:
- Cover page
- Table of contents (placeholder; Word will populate on open)
- Heading hierarchy
- Tables, bullets, callouts
- Mermaid code blocks (rendered as monospace; copy into any Mermaid renderer)
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


PROJECT = "Umang HIMS"
SUBTITLE = "AI-First Hospital Management System"
VERSION = "v1.0"
DOC_DATE = date(2026, 6, 1).strftime("%d %b %Y")
AUTHOR = "Solution Architecture Team"

PRIMARY = RGBColor(0x0F, 0x4C, 0x81)      # deep blue
ACCENT = RGBColor(0x29, 0x8A, 0xC8)        # light blue
INK = RGBColor(0x1A, 0x1F, 0x2C)           # near-black
MUTED = RGBColor(0x5F, 0x6B, 0x7A)         # grey
SUCCESS = RGBColor(0x1F, 0x7A, 0x3F)
WARN = RGBColor(0xB8, 0x6E, 0x00)
DANGER = RGBColor(0xB0, 0x2A, 0x37)


def _set_run(run, *, bold=False, italic=False, size=11, color=None, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def new_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    # Heading colours
    for lvl, sz in [("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12.5)]:
        try:
            s = styles[lvl]
            s.font.name = "Calibri"
            s.font.size = Pt(sz)
            s.font.color.rgb = PRIMARY
            s.font.bold = True
        except KeyError:
            pass

    return doc


def cover(doc: Document, doc_no: str, title: str, tagline: str) -> None:
    """Insert the standard cover page and a page break."""
    # Top spacer
    for _ in range(2):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    _set_run(r, bold=True, size=30, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    _set_run(r, italic=True, size=13, color=MUTED)

    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Document {doc_no}")
    _set_run(r, bold=True, size=14, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    _set_run(r, bold=True, size=26, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tagline)
    _set_run(r, italic=True, size=12, color=MUTED)

    for _ in range(8):
        doc.add_paragraph("")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    rows = [
        ("Version", VERSION),
        ("Date", DOC_DATE),
        ("Author", AUTHOR),
        ("Status", "Issued for review"),
    ]
    for i, (k, v) in enumerate(rows):
        c1, c2 = meta.rows[i].cells
        c1.text = ""
        c2.text = ""
        r = c1.paragraphs[0].add_run(k)
        _set_run(r, bold=True, color=PRIMARY)
        r = c2.paragraphs[0].add_run(v)
        _set_run(r, color=INK)

    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential — internal architecture & planning artefact")
    _set_run(r, italic=True, size=9, color=MUTED)

    doc.add_page_break()


def toc(doc: Document) -> None:
    """Insert a Word TOC field (populated when user presses F9 / opens in Word)."""
    h1(doc, "Contents")
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run._element.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_sep)

    # Placeholder text shown until the user updates the field in Word
    note = run.add_text("Right-click → Update Field to populate the table of contents.")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)

    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def lead(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    r = para.add_run(text)
    _set_run(r, italic=True, color=MUTED, size=11)
    para.paragraph_format.space_after = Pt(10)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    para = doc.add_paragraph(text, style="List Bullet")
    para.paragraph_format.left_indent = Cm(0.75 + 0.5 * level)
    para.paragraph_format.space_after = Pt(2)


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def kv(doc: Document, key: str, value: str) -> None:
    para = doc.add_paragraph()
    r = para.add_run(f"{key}: ")
    _set_run(r, bold=True, color=PRIMARY)
    r2 = para.add_run(value)
    _set_run(r2)


def callout(doc: Document, label: str, body: str, kind: str = "note") -> None:
    color = {"note": ACCENT, "warn": WARN, "danger": DANGER, "ok": SUCCESS}.get(kind, ACCENT)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    para = cell.paragraphs[0]
    r = para.add_run(f"{label}: ")
    _set_run(r, bold=True, color=color)
    r2 = para.add_run(body)
    _set_run(r2)
    # subtle shading
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F4F7FB")
    cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph("")


def table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_run(run, bold=True, color=PRIMARY)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run(run, size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_cm:
        for row in t.rows:
            for ci, w in enumerate(col_widths_cm):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph("")


def code_block(doc: Document, text: str, *, label: str | None = None) -> None:
    """Render a monospace block. We don't try to pretty-print syntax."""
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        _set_run(r, bold=True, italic=True, color=MUTED, size=9)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F6F6F6")
    cell._tc.get_or_add_tcPr().append(shading)
    para = cell.paragraphs[0]
    for line in text.splitlines() or [text]:
        run = para.add_run(line + "\n")
        _set_run(run, size=9.5, font="Consolas")
    doc.add_paragraph("")


def mermaid(doc: Document, caption: str, body: str) -> None:
    code_block(
        doc,
        body.strip("\n"),
        label=f"Mermaid diagram — {caption} (paste into any Mermaid renderer)",
    )


def page_break(doc: Document) -> None:
    doc.add_page_break()


def save(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  saved -> {path.name}")
